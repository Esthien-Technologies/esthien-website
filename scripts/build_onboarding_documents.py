from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "company-documents"
BRAND = ROOT / "public" / "brand-kit"

CORE = RGBColor(11, 40, 48)
BLUE = RGBColor(31, 116, 181)
MUTED = RGBColor(74, 92, 99)
LIGHT = "F2F4F7"

LEGAL_NOTE = (
    "Draft onboarding template only. This document is not legal advice and must be reviewed "
    "by India-qualified counsel/company secretary/tax advisor before use, signature, filing, "
    "or issuance."
)


def set_font(run, size=None, color=None, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph_border_bottom(paragraph, color="CBD5D8", size="4", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=90, start=130, bottom=90, end=130):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tbl_cell_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
    set_cell_margins(table)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, CORE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mark = BRAND / "esthien-mark.png"
    if mark.exists():
        p.add_run().add_picture(str(mark), width=Inches(0.24))
        p.add_run("  ")
    run = p.add_run("ESTHIEN LABS")
    set_font(run, 9, CORE, True)
    run = p.add_run("  |  Esthien Pvt Ltd  |  Onboarding Documentation")
    set_font(run, 9, MUTED)
    paragraph_border_bottom(p)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Esthien Pvt Ltd - confidential onboarding template - review before use - contact@esthien.com")
    set_font(run, 8.5, MUTED)


def add_p(doc, text, size=11, color=None, bold=False, style=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_font(run, size=size, color=color or RGBColor(0, 0, 0), bold=bold)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_font(run, 10.5, RGBColor(0, 0, 0))
    return p


def add_table(doc, rows, widths=(1900, 7460)):
    table = doc.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    set_table_geometry(table, list(widths))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            if c_idx == 0:
                shade_cell(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, 10, CORE if c_idx == 0 else RGBColor(0, 0, 0), c_idx == 0)
    return table


def signature_block(doc):
    add_heading(doc, "Execution and Acknowledgement", 2)
    add_table(
        doc,
        [
            ("Company", "Esthien Pvt Ltd / Esthien Labs [confirm exact legal name and CIN]"),
            ("Recipient", "[name, role, address, ID/passport if required]"),
            ("Signature", "[name, date, place]"),
            ("Review note", LEGAL_NOTE),
        ],
    )


DOCUMENTS = [
    {
        "title": "1. Onboarding Checklist and Document Acknowledgement",
        "purpose": "Create one intake record for interns, employees, consultants, advisors, and technical contributors.",
        "sections": [
            ("Required Before Start", [
                "Signed offer/internship agreement or engagement agreement.",
                "Signed NDA and PIIAA before access to source code, designs, prototypes, datasets, credentials, or partner information.",
                "Identity, address, emergency contact, bank/payment, tax, and institute/college details where applicable.",
                "Company email, device, repository, password manager, and communication-channel access approved by an owner.",
            ]),
            ("Acknowledgement", [
                "Recipient confirms receipt of the onboarding pack, agrees to follow company policies, and understands that confidential information and inventions remain protected after the engagement ends.",
            ]),
        ],
    },
    {
        "title": "2. Internship Agreement",
        "purpose": "Standard internship terms for technical, research, design, operations, or business interns.",
        "sections": [
            ("Appointment", [
                "Role: [internship title]. Department/project: [insert]. Supervisor: [insert]. Start date: [insert]. End date: [insert]. Location/remote terms: [insert].",
                "The internship is for learning and project contribution. It does not guarantee employment, continued engagement, equity, or conversion.",
            ]),
            ("Stipend and Benefits", [
                "Stipend: INR [amount] per [month/project], subject to applicable deductions and required documentation.",
                "Reimbursements require prior written approval and valid receipts.",
            ]),
            ("Work Rules", [
                "Intern must follow attendance, reporting, confidentiality, security, lab, prototype, data, open-source, and communication policies.",
                "Intern may not represent the company publicly, contact partners/customers/investors, publish work, or submit code/designs externally without written approval.",
            ]),
            ("IP and Confidentiality", [
                "The NDA and PIIAA are incorporated into this internship agreement. Work product created during the internship or using company resources belongs to the company to the fullest extent permitted by law.",
            ]),
            ("Termination", [
                "Either party may end the internship with [notice period] written notice unless misconduct, confidentiality breach, security breach, or policy violation requires immediate termination.",
            ]),
        ],
    },
    {
        "title": "3. Non-Disclosure Agreement (NDA)",
        "purpose": "Protect technical, commercial, investor, partner, prototype, and company information.",
        "sections": [
            ("Confidential Information", [
                "Includes source code, firmware, FPGA logic, schematics, PCB files, CAD, chipset architecture, prototypes, lab notes, roadmaps, datasets, biosignal/radar data, investor materials, customer/partner conversations, credentials, pricing, strategy, and non-public company operations.",
            ]),
            ("Obligations", [
                "Recipient will use confidential information only for authorized company work, protect it with reasonable care, not disclose it to third parties, and immediately report loss, unauthorized access, or suspected misuse.",
                "Recipient will not reverse engineer prototypes, copy repositories, export designs, train external models on company data, or upload confidential information to public tools without written approval.",
            ]),
            ("Exclusions", [
                "Information is not confidential if it is public through no breach, independently developed without confidential information, lawfully received from a third party, or approved for release in writing.",
            ]),
            ("Return and Survival", [
                "Upon request or exit, recipient will return or delete company materials. Confidentiality obligations survive the end of the relationship.",
            ]),
        ],
    },
    {
        "title": "4. Proprietary Information and Inventions Assignment Agreement (PIIAA)",
        "purpose": "Assign company work product and protect pre-existing inventions.",
        "sections": [
            ("Assignment", [
                "Recipient assigns to the company all rights in inventions, discoveries, designs, code, firmware, FPGA logic, chip architecture, schematics, documentation, models, datasets, processes, improvements, trade secrets, and works of authorship created within the scope of work or using company resources.",
                "Recipient will execute further documents required to perfect, register, enforce, or transfer these rights.",
            ]),
            ("Prior Inventions", [
                "Recipient must list all pre-existing inventions or third-party obligations before starting work. Unlisted prior inventions may be presumed not excluded from assignment where legally permitted.",
            ]),
            ("Moral Rights and Assistance", [
                "Recipient waives, assigns, or consents to company use of moral rights to the maximum extent permitted and will assist with patent, design, copyright, and trade-secret protection.",
            ]),
            ("Open Source and Third-Party Materials", [
                "Recipient may not introduce open-source code, third-party IP, datasets, design files, or AI-generated material into company work without review and written approval.",
            ]),
        ],
    },
    {
        "title": "5. Information Security and Acceptable Use Policy",
        "purpose": "Set baseline security rules for company systems and technical work.",
        "sections": [
            ("Access", [
                "Use unique passwords, approved password managers, MFA where available, least-privilege access, and named user accounts only.",
                "Do not share credentials, tokens, repository access, design files, prototypes, or company devices.",
            ]),
            ("Systems", [
                "Company work must stay in approved repositories, drives, communication tools, and issue trackers. Personal cloud drives and public AI tools require explicit approval for any company material.",
            ]),
            ("Incidents", [
                "Report lost devices, exposed credentials, suspicious messages, unauthorized repository access, data leaks, malware, or accidental sharing immediately.",
            ]),
        ],
    },
    {
        "title": "6. Open Source, AI Tools, and Third-Party Materials Policy",
        "purpose": "Prevent accidental IP contamination in semiconductor, robotics, and embedded-system work.",
        "sections": [
            ("Approval Required", [
                "Do not add third-party code, HDL, firmware, PCB/CAD files, datasets, vendor SDKs, model weights, generated code, or design references without owner review.",
                "Track package name, source URL, license, version, use case, and approval owner.",
            ]),
            ("Restricted Materials", [
                "Avoid copyleft or viral-license components in deliverables unless counsel/engineering leadership approves the exact use.",
                "Do not upload confidential company code, design files, datasets, or partner/customer data to external AI tools unless approved.",
            ]),
        ],
    },
    {
        "title": "7. Prototype, Lab, and Device Safety Policy",
        "purpose": "Set rules for hardware, medical-assistive prototypes, vehicle sensing, and lab materials.",
        "sections": [
            ("Prototype Handling", [
                "Only authorized personnel may handle powered prototypes, motor drivers, batteries, radar modules, sensors, prosthetic assemblies, or test rigs.",
                "No human testing, vehicle testing, clinical claims, or safety-critical demonstration may occur without written approval, risk review, and required supervision.",
            ]),
            ("Documentation", [
                "Record test setup, firmware/bitstream version, sensor configuration, observed issue, and corrective action. Photograph or export logs where appropriate.",
            ]),
        ],
    },
    {
        "title": "8. Data Protection and Privacy Policy",
        "purpose": "Protect personal, technical, medical, radar, partner, and company data.",
        "sections": [
            ("Data Handling", [
                "Collect only necessary data, store it in approved locations, limit access, use retention periods, and delete or anonymize data when no longer needed.",
                "Sensitive categories include biosignals, medical-assistive test data, radar logs, video/images, IDs, employee/intern details, investor information, source code, credentials, and security logs.",
            ]),
            ("Processing Rules", [
                "Do not transfer personal or sensitive data externally, to vendors, to AI tools, or across jurisdictions without approval and required processing terms.",
            ]),
        ],
    },
    {
        "title": "9. Code of Conduct, POSH, and Workplace Conduct Policy",
        "purpose": "Set expected conduct for employees, interns, consultants, advisors, and visitors.",
        "sections": [
            ("Conduct", [
                "Treat colleagues, interns, visitors, vendors, founders, and partners with respect. Harassment, discrimination, retaliation, intimidation, bullying, or unsafe conduct is prohibited.",
            ]),
            ("Reporting", [
                "Concerns may be reported to [founder/HR/contact]. POSH and grievance processes must follow applicable Indian law and company procedure once formally adopted.",
            ]),
        ],
    },
    {
        "title": "10. Exit, Return of Property, and Continuing Obligations",
        "purpose": "Close access and preserve confidentiality/IP duties after engagement ends.",
        "sections": [
            ("Exit Checklist", [
                "Return devices, prototypes, access cards, documents, lab materials, samples, credentials, and company property.",
                "Confirm deletion/return of local copies, exported repositories, design files, datasets, notes, and confidential information.",
            ]),
            ("Continuing Duties", [
                "Confidentiality, IP assignment, non-use of company materials, non-disparagement where enforceable, and cooperation obligations continue after exit.",
            ]),
        ],
    },
]


def add_cover(doc):
    add_p(doc, "ESTHIEN LABS", bold=True, color=CORE, size=12, after=2)
    add_p(doc, "(ESTHIEN PVT LTD)", bold=True, color=MUTED, size=10, after=8)
    add_p(doc, "Standard Onboarding Documentation Pack", bold=True, color=CORE, size=24, before=16, after=6)
    add_p(doc, "Internship agreement, NDA, PIIAA, and operating policies for technical contributors.", color=MUTED, size=13, after=16)
    add_table(
        doc,
        [
            ("Public brand", "Esthien Labs"),
            ("Registered name", "Esthien Pvt Ltd [confirm exact legal name and CIN]"),
            ("Use", "Interns, employees, consultants, advisors, collaborators, and technical contributors"),
            ("Prepared", "1 July 2026"),
        ],
    )
    add_p(doc, LEGAL_NOTE, size=11, color=RGBColor(156, 28, 28), bold=True, before=20, after=10)


def add_document(doc, item, index):
    if index > 1:
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_header_footer(doc.sections[-1])
    add_heading(doc, item["title"], 1)
    add_table(doc, [("Purpose", item["purpose"]), ("Status", "Draft template - legal review required before use")])
    for title, bullets in item["sections"]:
        add_heading(doc, title, 2)
        for bullet in bullets:
            add_bullet(doc, bullet)
    signature_block(doc)


def build_pack():
    OUT.mkdir(exist_ok=True)
    doc = Document()
    configure_doc(doc)
    set_header_footer(doc.sections[0])
    add_cover(doc)
    doc.add_page_break()
    add_heading(doc, "Document Index", 1)
    for item in DOCUMENTS:
        add_bullet(doc, f"{item['title']} - {item['purpose']}")
    for index, item in enumerate(DOCUMENTS, start=1):
        add_document(doc, item, index)
    path = OUT / "Esthien_Standard_Onboarding_Documentation_Pack.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_pack())
