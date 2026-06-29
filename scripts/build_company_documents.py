from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "company-documents"
BRAND = ROOT / "public" / "brand-kit"

CORE = RGBColor(10, 13, 16)
PANEL = RGBColor(17, 24, 28)
CYAN = RGBColor(76, 201, 240)
GREEN = RGBColor(50, 214, 154)
LINE = RGBColor(203, 213, 216)
MUTED = RGBColor(74, 92, 99)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(248, 250, 249)

LEGAL_NOTE = (
    "Draft template only. This document is not legal advice, is not a statutory filing, "
    "and should be reviewed by an India-qualified advocate/company secretary/tax advisor "
    "before signature, issuance, board approval, or filing."
)

SOURCES = [
    (
        "Companies Act, 2013 - Section 173",
        "India Code record for Board meetings; includes board meeting notice requirements.",
        "https://www.indiacode.nic.in/show-data?actid=AC_CEN_22_29_00008_201318_1517807327856&orderno=177&sectionId=49099&sectionno=173",
    ),
    (
        "Secretarial Standard on Board Meetings (SS-1)",
        "ICSI standard on Board meeting practice, effective from 1 April 2024.",
        "https://www.icsi.edu/media/webmodules/SS-1_1_2024.pdf",
    ),
    (
        "Companies Act, 2013 - Section 62",
        "India Code record for further issue of share capital, including employee stock option schemes.",
        "https://www.indiacode.nic.in/show-data?actid=AC_CEN_22_29_00008_201318_1517807327856&orderno=64&sectionId=1252&sectionno=62",
    ),
    (
        "Sexual Harassment of Women at Workplace Act, 2013",
        "India Code PDF for POSH policy, grievance, and internal committee alignment.",
        "https://www.indiacode.nic.in/bitstream/123456789/2104/1/A2013-14.pdf",
    ),
    (
        "Digital Personal Data Protection Act, 2023",
        "Official MeitY PDF for digital personal data processing duties and privacy policy alignment.",
        "https://www.meity.gov.in/static/uploads/2024/06/2bf1f0e9f04e6fb4f8fef35e82c42aa5.pdf",
    ),
]


def set_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph_border_bottom(paragraph, color="CBD5D8", size="8", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        elem = margins.find(qn(f"w:{name}"))
        if elem is None:
            elem = OxmlElement(f"w:{name}")
            margins.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def add_p(doc, text="", style=None, bold=False, color=None, size=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text: str):
    p = add_p(doc, style="List Bullet", after=3)
    run = p.add_run(text)
    set_font(run, 10.5, BLACK)
    return p


def add_number(doc, text: str):
    p = add_p(doc, style="List Number", after=3)
    run = p.add_run(text)
    set_font(run, 10.5, BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
      size, color, before, after = 16, RGBColor(46, 116, 181), 16, 8
    elif level == 2:
      size, color, before, after = 13, RGBColor(46, 116, 181), 12, 6
    else:
      size, color, before, after = 12, RGBColor(31, 77, 120), 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for run in p.runs:
        set_font(run, size=size, color=color, bold=True)
    return p


def add_label_table(doc, rows: list[tuple[str, str]], label_width=2160, value_width=7200):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [label_width, value_width])
    set_cell_margins(table)
    for r_idx, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[r_idx].cells
        if r_idx == 0:
            set_cell_shading(label_cell, "F2F4F7")
        label_cell.text = ""
        value_cell.text = ""
        p1 = label_cell.paragraphs[0]
        run = p1.add_run(label)
        set_font(run, 10, CORE, True)
        p2 = value_cell.paragraphs[0]
        run = p2.add_run(value)
        set_font(run, 10, BLACK)
    doc.add_paragraph()
    return table


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_section_brand_header_footer(section, label: str):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if (BRAND / "esthien-mark.png").exists():
        run = p.add_run()
        run.add_picture(str(BRAND / "esthien-mark.png"), width=Inches(0.28))
        p.add_run("  ")
    run = p.add_run("ESTHIEN LABS")
    set_font(run, 9, CORE, True)
    run = p.add_run(f"  |  {label}")
    set_font(run, 9, MUTED)
    paragraph_border_bottom(p, "CBD5D8", "4", "4")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Confidential template - review before use - contact@esthien.com")
    set_font(run, 8.5, MUTED)


def set_brand_header_footer(doc: Document, label: str):
    for section in doc.sections:
        set_section_brand_header_footer(section, label)


def add_cover(doc: Document, title: str, subtitle: str, classification: str):
    add_p(doc, "ESTHIEN LABS", bold=True, color=CORE, size=12, after=2)
    p = add_p(doc, title, bold=True, color=CORE, size=24, after=6)
    p.paragraph_format.space_before = Pt(18)
    add_p(doc, subtitle, color=MUTED, size=13, after=16)
    add_label_table(
        doc,
        [
            ("Company", "Esthien Labs / Esthien Labs Private Limited [confirm exact legal name and CIN]"),
            ("Classification", classification),
            ("Use", "Internal drafting, legal review, HR setup, board review, investor preparation"),
            ("Date", "Prepared on 29 June 2026"),
            ("Owner", "Founder / Company Secretary / Legal Counsel"),
        ],
    )
    note = add_p(doc, LEGAL_NOTE, bold=True, color=RGBColor(155, 28, 28), size=10.5, after=16)
    paragraph_border_bottom(note, "9B1C1C", "6", "8")
    add_p(
        doc,
        "This pack assumes an India-based private limited company. Replace every bracketed placeholder, confirm applicable state labour rules, tax treatment, sector-specific permissions, foreign investment conditions, and board/shareholder approval routes before use.",
        size=10.5,
        color=BLACK,
    )
    doc.add_page_break()


def add_sources_section(doc: Document):
    add_heading(doc, "Legal Basis and Review Notes", 1)
    add_p(
        doc,
        "These templates were structured around common India private limited company workflows. The pack references public legal materials for orientation only; counsel must confirm current applicability before execution.",
        size=10.5,
    )
    for title, desc, url in SOURCES:
        add_bullet(doc, f"{title}: {desc} {url}")
    add_p(
        doc,
        "Common governance reminders to verify: first board meeting timeline, board notice period, annual board meeting cadence, statutory registers, share issuance authority, ESOP approval route, employment law classification, stamp duty, tax withholding, privacy notices, and sector-specific medical/automotive compliance.",
        bold=True,
        color=CORE,
        size=10.5,
        after=12,
    )


def template_paragraphs(sections: list[tuple[str, Iterable[str]]]):
    return sections


TEMPLATES = [
    {
        "title": "01. Incorporation and Statutory Setup Checklist",
        "category": "Corporate Governance",
        "purpose": "Create the first operating record after incorporation and prevent missing statutory basics.",
        "sections": template_paragraphs([
            ("Template", [
                "Company legal name: [Esthien Labs Private Limited]. CIN: [insert]. Registered office: [insert].",
                "Prepare and maintain statutory registers, certificate of incorporation, PAN/TAN/GST records if applicable, common seal policy if adopted, digital signature records, board file, shareholder file, contracts file, employment file, and IP assignment file.",
                "Confirm first board meeting date and first set of board resolutions. Maintain signed attendance, notice, agenda, minutes, and resolutions in one indexed folder.",
            ]),
            ("Approval and Records", [
                "Owner: Founder / Company Secretary.",
                "Review cadence: monthly until all incorporation actions are closed, then quarterly.",
                "Evidence to store: filed forms, challans, board minutes, registers, bank KYC, tax registration confirmations, domain and email ownership evidence.",
            ]),
        ]),
    },
    {
        "title": "02. Consent to Act as Director and Disclosure Pack",
        "category": "Corporate Governance",
        "purpose": "Capture director consent, identity, interests, and ongoing disclosure obligations.",
        "sections": template_paragraphs([
            ("Template", [
                "I, [director name], having DIN/PAN/passport [insert], consent to act as Director of [company]. I confirm that I am not disqualified from appointment and will disclose interests, changes in particulars, and conflicts as required.",
                "I agree to act in good faith, exercise reasonable care, maintain confidentiality, and preserve company records received in my capacity as director.",
                "Attachments: identity proof, address proof, DIN details, consent form, disclosure of interest, list of other directorships/shareholdings, and contact details for board notices.",
            ]),
            ("Signature", [
                "Signed by: [director]. Date: [insert]. Place: [insert]. Witness: [insert if required].",
            ]),
        ]),
    },
    {
        "title": "03. First Board Meeting Notice and Agenda",
        "category": "Corporate Governance",
        "purpose": "Provide a reusable first board meeting notice and agenda.",
        "sections": template_paragraphs([
            ("Notice", [
                "Notice is hereby given that the first meeting of the Board of Directors of [company] will be held on [date] at [time] at [venue/video link], to transact the agenda below.",
                "Agenda: take note of incorporation, appoint chairperson, approve registered office records, take note of subscribers and share capital, approve bank account opening, adopt statutory registers, authorize filings, approve brand/domain/email assets, appoint advisors, and approve initial policies.",
            ]),
            ("Dispatch Record", [
                "Notice sent to: [director list]. Mode: [email/courier/hand delivery]. Date and time: [insert]. Proof retained: [insert].",
            ]),
        ]),
    },
    {
        "title": "04. First Board Meeting Minutes and Resolutions",
        "category": "Corporate Governance",
        "purpose": "Record foundational board approvals in a defensible minutes format.",
        "sections": template_paragraphs([
            ("Minutes Template", [
                "Meeting of the Board of Directors of [company] held on [date] at [time]. Directors present: [insert]. Invitees: [insert]. Chair: [insert].",
                "The Board took note of incorporation, certificate of incorporation, registered office, subscribers, initial shareholding, statutory records, and commencement actions.",
                "Resolved that [authorized person] is authorized to open and operate bank accounts, execute KYC, appoint tax/accounting/legal vendors, maintain statutory registers, approve company letterhead and brand assets, and perform all incidental acts.",
            ]),
            ("Close", [
                "There being no further business, the meeting concluded at [time] with a vote of thanks to the Chair.",
            ]),
        ]),
    },
    {
        "title": "05. Shareholder Written Consent / EGM Notice and Minutes",
        "category": "Corporate Governance",
        "purpose": "Provide shareholder approval formats for matters requiring member consent.",
        "sections": template_paragraphs([
            ("Template", [
                "The members of [company] approve the following resolution as [ordinary/special] resolution: [full resolution text].",
                "Background and explanatory note: [describe purpose, legal basis, valuation/tax consideration, effect on rights, filings required, and interested parties].",
                "Voting record: members present or consenting, shareholding, votes for/against/abstain, and authorization for filings.",
            ]),
            ("Records", [
                "Maintain notice, explanatory statement, attendance/proxy if applicable, signed minutes, shareholder consent, filings, and updated registers.",
            ]),
        ]),
    },
    {
        "title": "06. Statutory Register and Records Index",
        "category": "Corporate Governance",
        "purpose": "Create a practical index for registers, minute books, and approvals.",
        "sections": template_paragraphs([
            ("Index", [
                "Register of members; register of directors and key managerial personnel; register of charges; register of contracts and arrangements in which directors are interested; share certificates; minutes of board meetings; minutes of general meetings; filings and challans; contracts and IP assignments.",
                "Each register should record version owner, storage location, backup method, last review date, and access control.",
            ]),
            ("Control", [
                "Do not alter executed records without a dated correction note and approval route. Preserve scan and original separately.",
            ]),
        ]),
    },
    {
        "title": "07. Board Resolution for Bank Account and Finance Authority",
        "category": "Corporate Governance",
        "purpose": "Authorize bank account opening, signatories, and transaction controls.",
        "sections": template_paragraphs([
            ("Resolution", [
                "Resolved that a current account be opened with [bank] in the name of [company], and that [names/designations] are authorized signatories.",
                "Transaction authority: up to INR [amount] by [single/joint] authorization; above INR [amount] requires board/founder approval. Online banking, cards, UPI, and payment gateways are authorized subject to internal controls.",
            ]),
            ("Controls", [
                "Monthly bank reconciliation, expense documentation, maker-checker access, vendor onboarding approval, and no personal expenses without written reimbursement policy.",
            ]),
        ]),
    },
    {
        "title": "08. Board Resolution for Brand, Domain, Letterhead, and Public Profiles",
        "category": "Corporate Governance",
        "purpose": "Approve official identity assets and public communication channels.",
        "sections": template_paragraphs([
            ("Resolution", [
                "Resolved that the company adopt the Esthien Labs brand system, domain [www.esthien.com], official email addresses, logo assets, public social profiles, and letterhead templates as official communication assets.",
                "Authorized persons may maintain DNS, hosting, email, website, LinkedIn, Instagram, and related profiles, subject to brand, privacy, and security controls.",
            ]),
            ("Records", [
                "Store registrar ownership, admin access list, 2FA status, logo files, brand kit, and public profile URLs in the corporate records folder.",
            ]),
        ]),
    },
    {
        "title": "09. Founder Agreement",
        "category": "Founder Documents",
        "purpose": "Set founder roles, equity expectations, governance, and departure handling.",
        "sections": template_paragraphs([
            ("Core Terms", [
                "Parties: [founder names] and [company]. Effective date: [insert]. Purpose: build Esthien Labs as a deep-tech company focused on FPGA chipsets, physical intelligence, medical assistive systems, and automotive sensing.",
                "Roles: each founder will perform duties assigned by the Board and will act in the best interests of the company.",
                "Equity: initial shareholding, vesting/reverse vesting, cliff, acceleration, repurchase rights, transfer restrictions, and treatment of bad leaver/good leaver events must be finalized with counsel.",
                "IP: all work product, inventions, designs, source code, technical documents, brand assets, domain assets, and business materials created for the company belong to the company.",
                "Disputes: escalation to board discussion, mediation, and jurisdiction/venue [insert].",
            ]),
            ("Signatures", [
                "Founder 1: [insert]. Founder 2: [insert]. Company: [authorized signatory].",
            ]),
        ]),
    },
    {
        "title": "10. Founder Intellectual Property Assignment",
        "category": "Founder Documents",
        "purpose": "Assign pre-incorporation and ongoing founder-created IP to the company.",
        "sections": template_paragraphs([
            ("Assignment", [
                "Assignor irrevocably assigns to [company] all rights, title, and interest in company-related inventions, designs, software, documentation, schematics, chip architecture concepts, brand assets, domain assets, trade secrets, research notes, and business materials created before or after incorporation for company purposes.",
                "Assignor will execute further documents and assist with filings, registrations, patent applications, copyright records, and enforcement.",
                "Excluded IP must be listed in Schedule A. Any unlisted company-related material is presumed assigned to the company.",
            ]),
            ("Moral Rights and Confidentiality", [
                "To the extent permitted by law, assignor waives or agrees not to assert moral rights against company use and will keep all non-public information confidential.",
            ]),
        ]),
    },
    {
        "title": "11. Founder Confidentiality and Non-Disclosure Agreement",
        "category": "Founder Documents",
        "purpose": "Protect sensitive founder-stage information.",
        "sections": template_paragraphs([
            ("Confidential Information", [
                "Includes technical architectures, FPGA designs, chipset roadmaps, medical/automotive concepts, vendor terms, investor materials, customer discussions, financial models, source code, business plans, and internal communications.",
                "Recipient may use confidential information only for company purposes and must protect it using reasonable safeguards.",
            ]),
            ("Exceptions and Term", [
                "Exceptions: information already public through no breach, independently developed without use of confidential information, or lawfully received from a third party.",
                "Confidentiality continues until information becomes public through no breach, with trade secrets protected as long as they remain trade secrets.",
            ]),
        ]),
    },
    {
        "title": "12. Founder Service / Employment Agreement",
        "category": "Founder Documents",
        "purpose": "Set founder duties, compensation, confidentiality, IP, and termination basics.",
        "sections": template_paragraphs([
            ("Terms", [
                "Position: [Founder/CEO/CTO]. Reporting: Board of Directors. Start date: [insert]. Compensation: INR [insert] per month or deferred until board approval.",
                "Duties include strategy, hiring, product, technology, fundraising, partnerships, compliance, and any duties assigned by the Board.",
                "Founder will comply with company policies, protect confidential information, assign company IP, avoid conflicts, and return company property on termination.",
            ]),
            ("Termination", [
                "Termination, notice, garden leave, severance, equity treatment, and access removal should align with the founder agreement and board approvals.",
            ]),
        ]),
    },
    {
        "title": "13. Founder Vesting and Reverse Vesting Term Sheet",
        "category": "Founder Documents",
        "purpose": "Capture founder equity vesting principles before definitive agreements.",
        "sections": template_paragraphs([
            ("Term Sheet", [
                "Vesting schedule: [four years with one-year cliff / custom]. Commencement: [date]. Reverse vesting applies to founder shares subject to company repurchase on departure.",
                "Good leaver: [death, disability, termination without cause, board-approved reason]. Bad leaver: [fraud, misconduct, IP breach, confidentiality breach, competing business, resignation before cliff].",
                "Acceleration: [none/single-trigger/double-trigger] to be approved by shareholders/investors if applicable.",
            ]),
            ("Implementation", [
                "Definitive documents may include shareholders agreement, employment/service agreement, share transfer restrictions, Articles amendments, and board/shareholder approvals.",
            ]),
        ]),
    },
    {
        "title": "14. Conflict of Interest and Related Party Disclosure",
        "category": "Founder Documents",
        "purpose": "Capture founder/director conflicts and related party transactions.",
        "sections": template_paragraphs([
            ("Disclosure", [
                "Discloser: [name]. Role: [director/founder/employee]. I disclose the following direct or indirect interests, affiliations, investments, contracts, family relationships, or competing activities that may affect company decisions: [insert].",
                "I will update this disclosure promptly when circumstances change and will abstain from decisions where required.",
            ]),
            ("Board Action", [
                "Board review outcome: [approved/approved with conditions/not approved]. Conditions: [insert].",
            ]),
        ]),
    },
    {
        "title": "15. Investor Confidentiality Agreement",
        "category": "Investor Documents",
        "purpose": "Protect company information shared during fundraising or diligence.",
        "sections": template_paragraphs([
            ("Terms", [
                "Recipient will use confidential information solely to evaluate a potential investment, strategic partnership, or financing transaction with [company].",
                "Recipient may share information only with representatives who need to know and are bound by confidentiality duties.",
                "No license is granted to patents, designs, source code, trade secrets, brand assets, data, or other IP.",
            ]),
            ("Return and Remedies", [
                "On request, recipient will return or destroy confidential materials except archival/legal retention copies. Company may seek injunctive relief for breach.",
            ]),
        ]),
    },
    {
        "title": "16. Seed Investment Term Sheet",
        "category": "Investor Documents",
        "purpose": "Summarize key commercial terms before definitive investment documents.",
        "sections": template_paragraphs([
            ("Indicative Terms", [
                "Issuer: [company]. Investor: [name]. Instrument: [equity/CCPS/CCD/SAFE-like contractual instrument subject to Indian law review]. Investment amount: INR [insert]. Valuation/cap: [insert].",
                "Use of proceeds: product development, FPGA prototypes, team hiring, IP filings, compliance, pilots, and working capital.",
                "Rights: information rights, reserved matters, pro-rata rights, board observer/director rights, transfer restrictions, anti-dilution, liquidation preference, founder vesting, ESOP pool, and warranties.",
            ]),
            ("Status", [
                "Non-binding except confidentiality, exclusivity, costs, governing law, and dispute resolution if stated as binding.",
            ]),
        ]),
    },
    {
        "title": "17. Share Subscription Agreement - Key Template",
        "category": "Investor Documents",
        "purpose": "Provide a working skeleton for subscription of shares/securities.",
        "sections": template_paragraphs([
            ("Core Clauses", [
                "Investor subscribes for [number/class] securities at INR [price] per security, subject to conditions precedent, board/shareholder approvals, valuation report if applicable, filings, and receipt of funds.",
                "Company warranties: incorporation, authority, capitalization, IP ownership, no undisclosed debt, compliance, tax filings, litigation, employment, data protection, and contracts.",
                "Investor warranties: authority, funds, compliance with applicable laws, KYC, beneficial ownership, and investment intent.",
            ]),
            ("Completion", [
                "At closing: funds received, securities allotted, share certificates/records issued, registers updated, filings completed, and closing deliverables exchanged.",
            ]),
        ]),
    },
    {
        "title": "18. Shareholders Agreement - Key Template",
        "category": "Investor Documents",
        "purpose": "Set governance and economic rights between founders, investors, and company.",
        "sections": template_paragraphs([
            ("Governance", [
                "Board composition, quorum, reserved matters, information rights, budgets, audit, related party transactions, founder commitments, and compliance obligations.",
                "Transfer restrictions: right of first refusal, right of first offer, tag-along, drag-along, permitted transfers, lock-in, and prohibited transfers.",
                "Economic rights: liquidation preference, anti-dilution, conversion, dividends, exit rights, pre-emptive rights, and ESOP pool management.",
            ]),
            ("Company Protection", [
                "Founder IP assignment, confidentiality, non-solicit, conflict restrictions, data/security duties, and return of property.",
            ]),
        ]),
    },
    {
        "title": "19. Convertible Preference Share / Convertible Instrument Term Sheet",
        "category": "Investor Documents",
        "purpose": "Outline a possible India-reviewed convertible financing structure.",
        "sections": template_paragraphs([
            ("Terms", [
                "Instrument: [CCPS/CCD/other counsel-approved instrument]. Principal/subscription amount: INR [insert]. Conversion trigger: qualified financing, long-stop date, acquisition, IPO, or voluntary conversion.",
                "Conversion price: [valuation cap/discount/fixed price] subject to applicable law, valuation, pricing guidelines, tax, foreign investment rules, and Articles authorization.",
                "Investor protections: information rights, reserved matters, transfer rights, liquidation preference, anti-dilution, and dispute resolution.",
            ]),
            ("Compliance", [
                "Do not issue until counsel confirms Companies Act, FEMA/FDI, tax, valuation, stamp duty, and filing requirements.",
            ]),
        ]),
    },
    {
        "title": "20. ESOP Scheme and Grant Letter",
        "category": "Investor / Employee Equity",
        "purpose": "Create an employee stock option framework and grant template.",
        "sections": template_paragraphs([
            ("Scheme Terms", [
                "Pool size: [percentage/number]. Eligibility: employees/directors/consultants as permitted by law and scheme terms. Vesting: [schedule], with minimum vesting and approval requirements confirmed by counsel.",
                "Exercise price, exercise period, lapse, resignation, termination, death/disability, transfer restrictions, tax withholding, administration, disclosures, and register maintenance must be specified.",
            ]),
            ("Grant Letter", [
                "Grantee: [name]. Options: [number]. Grant date: [date]. Vesting commencement: [date]. Exercise price: INR [insert]. Acceptance deadline: [date].",
                "This grant is subject to the ESOP scheme, board/shareholder approvals, applicable law, and continued service.",
            ]),
        ]),
    },
    {
        "title": "21. Employee Offer Letter",
        "category": "Employee Documents",
        "purpose": "Issue a pre-joining employment offer.",
        "sections": template_paragraphs([
            ("Offer", [
                "Dear [candidate], we are pleased to offer you the position of [title] at [company], reporting to [manager], with proposed start date [date] and work location [location/remote].",
                "Compensation: fixed salary INR [insert] per annum, variable/bonus [insert], benefits [insert], equity eligibility [if any, subject to ESOP approvals].",
                "This offer is conditional on background verification, identity/right-to-work checks, execution of appointment letter, IP assignment, confidentiality agreement, and company policies.",
            ]),
            ("Acceptance", [
                "Please sign and return by [date]. This offer does not create employment until joining formalities are completed and the appointment letter is executed.",
            ]),
        ]),
    },
    {
        "title": "22. Appointment Letter / Employment Agreement",
        "category": "Employee Documents",
        "purpose": "Set core employment terms after acceptance.",
        "sections": template_paragraphs([
            ("Terms", [
                "Employee: [name]. Position: [title]. Start date: [date]. Probation: [period]. Work location: [insert]. Reporting manager: [insert]. Hours and leave: as per company policy and applicable law.",
                "Duties: perform assigned responsibilities, comply with policies, protect company property and confidential information, assign IP, avoid conflicts, and follow security procedures.",
                "Termination: notice period [insert], immediate termination for cause, return of property, final settlement, confidentiality survival, and post-employment restrictions where enforceable.",
            ]),
            ("Policies", [
                "Employee acknowledges receipt of handbook, IT/security policy, anti-harassment policy, privacy notice, expense policy, and IP/confidentiality obligations.",
            ]),
        ]),
    },
    {
        "title": "23. Employee Invention and IP Assignment Agreement",
        "category": "Employee Documents",
        "purpose": "Assign employee-created work product and inventions to the company.",
        "sections": template_paragraphs([
            ("Assignment", [
                "Employee assigns to [company] all work product, inventions, designs, code, schematics, documents, data, models, processes, trade secrets, and improvements created within employment scope or using company resources.",
                "Employee will promptly disclose inventions and assist with registrations, patent filings, prosecution, enforcement, and execution of further documents.",
                "Excluded prior inventions must be listed in Schedule A before employment begins.",
            ]),
            ("Open Source and Third-Party Materials", [
                "Employee may not introduce third-party code, datasets, designs, or open-source components into company work without approval and license review.",
            ]),
        ]),
    },
    {
        "title": "24. Employee Confidentiality and Data Security Agreement",
        "category": "Employee Documents",
        "purpose": "Protect technical, commercial, and personal data handled by employees.",
        "sections": template_paragraphs([
            ("Obligations", [
                "Employee will protect confidential information, use company systems securely, follow access controls, report incidents promptly, and not remove or disclose data except for authorized company purposes.",
                "Special categories: medical prototypes, biosignal data, radar data, embedded system designs, vendor terms, customer data, investor materials, source code, chip architecture documents, and security credentials.",
            ]),
            ("Return", [
                "On request or termination, employee will return devices, documents, credentials, notes, source material, and copies, and certify deletion from personal devices where required.",
            ]),
        ]),
    },
    {
        "title": "25. Employee Handbook Core Policies",
        "category": "Employee Documents",
        "purpose": "Provide a starter handbook with basic policies.",
        "sections": template_paragraphs([
            ("Policies", [
                "Code of conduct: integrity, lawful conduct, respectful workplace, confidentiality, accurate records, anti-bribery, no unauthorized public statements, and conflict disclosure.",
                "Workplace: attendance, remote work, leave, reimbursement, travel, device use, information security, acceptable use, open-source review, procurement, and expense approvals.",
                "People: equal opportunity, anti-harassment, grievance reporting, performance reviews, disciplinary process, and separation process.",
            ]),
            ("Acknowledgement", [
                "Employee acknowledges that policies may be amended and that local law, employment contract, and mandatory statutory rights prevail over handbook language.",
            ]),
        ]),
    },
    {
        "title": "26. Anti-Harassment, POSH, and Grievance Policy",
        "category": "Employee Documents",
        "purpose": "Create a baseline workplace conduct and complaint framework.",
        "sections": template_paragraphs([
            ("Policy", [
                "The company prohibits sexual harassment, discrimination, retaliation, bullying, intimidation, and abusive conduct in the workplace, online channels, events, travel, and remote work.",
                "Complaints may be made to [email/person]. The company will treat complaints sensitively, prohibit retaliation, conduct a fair inquiry, and take appropriate action.",
                "If statutory POSH committee requirements apply, the company will constitute and maintain the Internal Committee, notices, training, and records as required.",
            ]),
            ("Acknowledgement", [
                "All employees, consultants, interns, founders, directors, and visitors must comply with this policy.",
            ]),
        ]),
    },
    {
        "title": "27. Consultant / Advisor Agreement",
        "category": "Operations Documents",
        "purpose": "Engage contractors, advisors, and technical consultants cleanly.",
        "sections": template_paragraphs([
            ("Terms", [
                "Consultant: [name]. Services: [scope]. Term: [dates]. Fees: INR [amount]/[milestone]. Expenses: [approval process]. Relationship: independent contractor, not employee.",
                "Deliverables and work product are assigned to the company upon creation/payment as specified. Consultant must maintain confidentiality, avoid conflicts, and comply with company security rules.",
                "Consultant will not bind the company, hire staff, incur expenses, or make public statements without written authorization.",
            ]),
            ("Termination", [
                "Either party may terminate with [notice]. Company may terminate immediately for breach, misconduct, confidentiality/IP breach, or security risk.",
            ]),
        ]),
    },
    {
        "title": "28. Vendor / Services Agreement",
        "category": "Operations Documents",
        "purpose": "Create a basic procurement and vendor services template.",
        "sections": template_paragraphs([
            ("Terms", [
                "Vendor will provide [services/products] under [SOW/PO]. Fees, taxes, payment terms, milestones, acceptance criteria, warranties, support, and service levels are set out in Schedule A.",
                "Vendor must protect confidential information, comply with security requirements, process personal data only on instructions if applicable, and maintain required licenses.",
                "IP ownership: company-owned deliverables belong to company; vendor background IP remains vendor property but company receives necessary license.",
            ]),
            ("Risk", [
                "Include indemnity, limitation of liability, insurance, audit rights, termination rights, and transition assistance based on deal risk.",
            ]),
        ]),
    },
    {
        "title": "29. Website Terms of Use",
        "category": "Public / Website Documents",
        "purpose": "Set public website terms for visitors.",
        "sections": template_paragraphs([
            ("Terms", [
                "The website provides general information about Esthien Labs. No product, medical, automotive safety, investment, employment, or technical advice is provided through the website.",
                "Visitors may not misuse the site, probe security, copy brand assets without permission, scrape content excessively, submit unlawful material, or imply endorsement.",
                "All content, brand assets, logos, text, diagrams, and visual systems are owned by or licensed to the company.",
            ]),
            ("Disclaimers", [
                "The site is provided as-is. Company may change content, restrict access, or update terms. Governing law and jurisdiction: [insert].",
            ]),
        ]),
    },
    {
        "title": "30. Privacy Policy and Data Processing Notice",
        "category": "Public / Website Documents",
        "purpose": "Create a starter privacy notice for website and business inquiries.",
        "sections": template_paragraphs([
            ("Notice", [
                "Data collected: name, email, organization, inquiry content, social profile interactions, website analytics, security logs, and any information voluntarily provided.",
                "Purposes: respond to inquiries, manage partnerships, investor conversations, recruitment, security, legal compliance, analytics, and business operations.",
                "Rights and requests: contact [privacy@esthien.com/contact@esthien.com]. The company will respond as required by applicable law.",
            ]),
            ("Controls", [
                "Maintain data inventory, access controls, retention schedule, breach response process, vendor processing terms, and privacy policy review cadence.",
            ]),
        ]),
    },
]


def add_template(doc: Document, item: dict, idx: int):
    if idx > 1:
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_brand_header_footer(doc.sections[-1], "Legal Foundation Pack")
    add_heading(doc, item["title"], 1)
    add_label_table(
        doc,
        [
            ("Category", item["category"]),
            ("Purpose", item["purpose"]),
            ("Status", "Draft template - legal review required before use"),
        ],
        label_width=1800,
        value_width=7560,
    )
    for heading, paragraphs in item["sections"]:
        add_heading(doc, heading, 2)
        for para in paragraphs:
            add_p(doc, para, size=10.5)
    add_heading(doc, "Execution Block", 2)
    add_label_table(
        doc,
        [
            ("Approved by", "[Board / shareholders / authorized signatory / employee / counterparty]"),
            ("Signature", "[name, designation, date, place]"),
            ("Record owner", "[Founder / HR / Company Secretary / Legal / Finance]"),
            ("Review note", LEGAL_NOTE),
        ],
        label_width=1800,
        value_width=7560,
    )


def build_legal_pack():
    doc = Document()
    configure_doc(doc)
    set_brand_header_footer(doc, "Legal Foundation Pack")
    add_cover(
        doc,
        "Legal Foundation Pack",
        "30 root templates for founders, employees, investors, governance, and company operations.",
        "Confidential - Draft templates for review",
    )
    add_sources_section(doc)
    doc.add_page_break()
    add_heading(doc, "Template Index", 1)
    for item in TEMPLATES:
        add_bullet(doc, f"{item['title']} - {item['category']}")
    for idx, item in enumerate(TEMPLATES, start=1):
        add_template(doc, item, idx)
    path = OUT / "Esthien_Legal_Foundation_Pack_30_Templates.docx"
    doc.save(path)
    return path


def add_letterhead_header(doc: Document, title: str):
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if (BRAND / "esthien-mark.png").exists():
        run = p.add_run()
        run.add_picture(str(BRAND / "esthien-mark.png"), width=Inches(0.28))
        p.add_run("  ")
    run = p.add_run("ESTHIEN LABS")
    set_font(run, 9, BLACK, True)
    run = p.add_run("\n  CUSTOM FPGA CHIPSETS FOR PHYSICAL INTELLIGENCE")
    set_font(run, 6.8, MUTED)
    paragraph_border_bottom(p, "2B373D", "5", "6")
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Esthien Labs | www.esthien.com | contact@esthien.com | [Registered Office Address] | CIN: [insert]")
    set_font(run, 8, MUTED)


def add_letter_template(doc: Document, title: str, rows: list[tuple[str, str]], body: list[str]):
    if len(doc.paragraphs) > 1:
        doc.add_section(WD_SECTION.NEW_PAGE)
    add_letterhead_header(doc, title)
    add_p(doc, title, bold=True, color=CORE, size=18, after=8, before=12)
    add_label_table(doc, rows, label_width=1800, value_width=7560)
    for para in body:
        add_p(doc, para, size=10.5)
    add_p(doc, "Sincerely,", size=10.5, after=18)
    add_p(doc, "[Authorized Signatory]\nEsthien Labs", size=10.5)


def build_letterhead_pack():
    doc = Document()
    configure_doc(doc)
    templates = [
        (
            "Blank Official Letterhead",
            [("Date", "[insert]"), ("Reference", "[insert]"), ("To", "[recipient]")],
            [
                "[Body text begins here. Use this format for official letters, notices, certificates, cover letters, bank letters, and government/vendor correspondence.]",
                "[Replace placeholders, keep one subject per letter, and attach supporting documents where required.]",
            ],
        ),
        (
            "Offer Letter (OL)",
            [("Candidate", "[name]"), ("Position", "[title]"), ("Start Date", "[date]"), ("Compensation", "INR [amount] per annum")],
            [
                "We are pleased to offer you the position stated above, subject to satisfactory verification, execution of appointment documentation, confidentiality/IP assignment terms, and company policies.",
                "This offer is conditional and does not create employment until joining formalities are completed. Please sign below to acknowledge acceptance by [date].",
            ],
        ),
        (
            "Appointment Letter",
            [("Employee", "[name]"), ("Role", "[title]"), ("Manager", "[name]"), ("Location", "[location/remote]")],
            [
                "Your appointment is governed by the terms in the employment agreement, employee handbook, confidentiality agreement, IP assignment, and applicable company policies.",
                "You will perform duties assigned by the company, preserve confidentiality, comply with security requirements, and assign company-related work product to the company.",
            ],
        ),
        (
            "Official Business Letter",
            [("Recipient", "[name/company]"), ("Subject", "[subject]"), ("Reference", "[reference number]")],
            [
                "This letter records the company's official position on the subject above. Any commercial or legal commitment must be supported by an authorized agreement, purchase order, board approval, or written authorization as applicable.",
                "Please contact contact@esthien.com for clarifications.",
            ],
        ),
        (
            "Board / Founder Certificate",
            [("Certificate", "[type]"), ("Date", "[date]"), ("Issued To", "[recipient]")],
            [
                "This is to certify that [statement]. This certificate is issued based on company records available as of the date above and is subject to verification against statutory registers and board/shareholder approvals.",
                "This certificate should not be used for statutory filing unless reviewed by the company secretary or counsel.",
            ],
        ),
        (
            "Investor Communication",
            [("Investor", "[name]"), ("Round", "[seed/pre-seed/other]"), ("Confidentiality", "Confidential")],
            [
                "This communication is provided for discussion purposes only and does not constitute an offer to sell securities, investment advice, or a binding commitment unless executed definitive documents state otherwise.",
                "All non-public information remains confidential and may not be shared without written consent from the company.",
            ],
        ),
        (
            "Employment Confirmation / Experience Letter",
            [("Employee", "[name]"), ("Role", "[title]"), ("Period", "[start date] to [end/current]")],
            [
                "This letter confirms the employment details above based on company records. It does not waive any continuing confidentiality, IP, non-solicit, return-of-property, or other surviving obligations.",
                "Issued on request of the employee for [purpose].",
            ],
        ),
    ]
    for template in templates:
        add_letter_template(doc, *template)
    path = OUT / "Esthien_Letterhead_and_OL_Templates.docx"
    doc.save(path)
    return path


def main():
    OUT.mkdir(exist_ok=True)
    legal = build_legal_pack()
    letterhead = build_letterhead_pack()
    print(legal)
    print(letterhead)


if __name__ == "__main__":
    main()
